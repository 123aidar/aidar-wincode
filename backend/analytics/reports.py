"""
Report generation utilities — Word (.docx) and Excel (.xlsx).
All Word documents comply with: Times New Roman 14pt, 1.5 line spacing, justified.
"""
import io
from decimal import Decimal
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import (
    Font as XFont, PatternFill, Alignment as XAlignment,
    Border, Side,
)
from openpyxl.utils import get_column_letter


# ─── Formatting helpers ────────────────────────────────────────────────────────

TNR = 'Times New Roman'

_HEADER_FILL  = PatternFill('solid', fgColor='1E293B')
_SUBHEAD_FILL = PatternFill('solid', fgColor='334155')
_ALT_FILL     = PatternFill('solid', fgColor='F1F5F9')
_BORDER_SIDE  = Side(style='thin', color='CBD5E1')
_THIN_BORDER  = Border(
    left=_BORDER_SIDE, right=_BORDER_SIDE,
    top=_BORDER_SIDE,  bottom=_BORDER_SIDE,
)


def _fmt(value):
    if value is None:
        return '0,00 ₽'
    try:
        return f"{float(value):,.2f} ₽".replace(',', ' ')
    except Exception:
        return str(value)


def _fmt_date(d):
    return d.strftime('%d.%m.%Y') if hasattr(d, 'strftime') else str(d)


# ─── python-docx helpers ───────────────────────────────────────────────────────

def _spacing_xml(paragraph, line='360', after='120'):
    """Inject w:spacing: line=360 twips (1.5×), after=120 (6pt)."""
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    el = OxmlElement('w:spacing')
    el.set(qn('w:line'),     line)
    el.set(qn('w:lineRule'), 'auto')
    el.set(qn('w:after'),    after)
    pPr.append(el)


def _rfonts_xml(run):
    """Ensure every run uses Times New Roman for ASCII, hAnsi and CS."""
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(attr), TNR)
    rPr.insert(0, rf)


def _setup_doc(doc):
    """Apply document defaults: margins, Normal style."""
    style = doc.styles['Normal']
    style.font.name = TNR
    style.font.size = Pt(14)
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(1.5)


def _p(doc, text, bold=False, size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
       space_after='120', center=False):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    _spacing_xml(para, after=space_after)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    run = para.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(size)
    run.font.bold = bold
    _rfonts_xml(run)
    return para


def _h(doc, text, level=1):
    """Add a styled heading (not using Word's Heading styles to keep TNR)."""
    sizes = {1: 16, 2: 14, 3: 14}
    size  = sizes.get(level, 14)
    para  = doc.add_paragraph()
    _spacing_xml(para, line='360', after='120')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(size)
    run.font.bold = True
    _rfonts_xml(run)
    return para


def _divider(doc):
    para = doc.add_paragraph()
    _spacing_xml(para, after='60')
    run = para.add_run('─' * 80)
    run.font.name = TNR
    run.font.size = Pt(10)
    _rfonts_xml(run)


def _table(doc, headers, rows, col_widths=None):
    """Add a nicely formatted table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run  = para.add_run(h)
        run.font.name  = TNR
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _rfonts_xml(run)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1E293B')
        tcPr.append(shd)

    # Data rows
    for r, row_data in enumerate(rows):
        fill = 'F1F5F9' if r % 2 == 1 else 'FFFFFF'
        for c, val in enumerate(row_data):
            cell = tbl.rows[r + 1].cells[c]
            cell.text = ''
            para = cell.paragraphs[0]
            run  = para.add_run(str(val))
            run.font.name = TNR
            run.font.size = Pt(11)
            _rfonts_xml(run)
            is_num = isinstance(val, (int, float, Decimal))
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_num else WD_ALIGN_PARAGRAPH.LEFT
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl


# ─── Excel helpers ─────────────────────────────────────────────────────────────

def _xh(ws, row, values, bold=True, bg='1E293B', fg='FFFFFF', size=12):
    """Write a header row to an Excel sheet."""
    fill = PatternFill('solid', fgColor=bg)
    font = XFont(name=TNR, size=size, bold=bold, color=fg)
    for col, val in enumerate(values, 1):
        cell = ws.cell(row=row, column=col, value=val)
        cell.font      = font
        cell.fill      = fill
        cell.alignment = XAlignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border    = _THIN_BORDER


def _xr(ws, row, values, alt=False, bold=False):
    """Write a data row."""
    fill = PatternFill('solid', fgColor='F1F5F9') if alt else PatternFill('solid', fgColor='FFFFFF')
    for col, val in enumerate(values, 1):
        cell = ws.cell(row=row, column=col, value=val)
        cell.font      = XFont(name=TNR, size=11, bold=bold)
        cell.fill      = fill
        cell.border    = _THIN_BORDER
        if isinstance(val, (int, float, Decimal)):
            cell.alignment = XAlignment(horizontal='right')
            if isinstance(val, Decimal) or (isinstance(val, float) and '.' in str(val)):
                cell.number_format = '#,##0.00'
        else:
            cell.alignment = XAlignment(horizontal='left')


def _xcol_widths(ws, widths):
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w


# ══════════════════════════════════════════════════════════════════════════════
#  WORD REPORT 1 — НАЛОГОВЫЙ ОТЧЁТ
# ══════════════════════════════════════════════════════════════════════════════

def build_word_tax_report(data: dict) -> io.BytesIO:
    doc = Document()
    _setup_doc(doc)

    df  = _fmt_date(data['date_from'])
    dt  = _fmt_date(data['date_to'])
    rev = data.get('total_revenue', Decimal('0'))
    parts_cost = data.get('parts_cost', Decimal('0'))
    orders_cnt = data.get('orders_count', 0)
    payments   = data.get('payment_by_method', {})
    monthly    = data.get('monthly_data', [])
    services   = data.get('services_revenue', [])
    gen_date   = _fmt_date(date.today())

    # ── Title block ──────────────────────────────────────────────────────────
    _p(doc, 'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ', bold=True, center=True)
    _p(doc, '«ВинАвто»', bold=True, size=16, center=True)
    _p(doc, ' ', size=6, space_after='0')

    _p(doc, 'ОТЧЁТ О ДОХОДАХ И НАЛОГОВЫХ ОБЯЗАТЕЛЬСТВАХ', bold=True, size=15, center=True)
    _p(doc, f'за период с {df} по {dt}', bold=False, size=13, center=True)
    _p(doc, ' ', size=6, space_after='0')
    _divider(doc)

    # ── Реквизиты ────────────────────────────────────────────────────────────
    _h(doc, '1. РЕКВИЗИТЫ ОРГАНИЗАЦИИ', level=1)
    _p(doc, 'Полное наименование: Общество с ограниченной ответственностью «ВинАвто»')
    _p(doc, 'Сокращённое наименование: ООО «ВинАвто»')
    _p(doc, 'ИНН/КПП: 7700000000 / 770001001')
    _p(doc, 'ОГРН: 1027700000000')
    _p(doc, 'Юридический адрес: Российская Федерация, г. Москва, ул. Автосервисная, д. 1')
    _p(doc, 'Система налогообложения: Упрощённая система налогообложения (УСН), объект '
            '«Доходы» — ставка 6%')
    _p(doc, f'Дата составления отчёта: {gen_date}')
    _p(doc, ' ', size=4, space_after='0')

    # ── Основные показатели ───────────────────────────────────────────────────
    _h(doc, '2. ОСНОВНЫЕ ПОКАЗАТЕЛИ ДЕЯТЕЛЬНОСТИ ЗА ПЕРИОД', level=1)
    _p(doc,
       f'За отчётный период с {df} по {dt} организацией ООО «ВинАвто» '
       f'была осуществлена хозяйственная деятельность по оказанию услуг технического '
       f'обслуживания и ремонта транспортных средств. Ниже представлены сводные '
       f'финансовые показатели за указанный период.')

    vat_rate   = Decimal('0.20')
    revenue_wo_vat = rev / (1 + vat_rate) if rev else Decimal('0')
    vat_amount = rev - revenue_wo_vat

    usn_tax = rev * Decimal('0.06')
    usn_tax_15 = max((rev - parts_cost) * Decimal('0.15'), Decimal('0'))

    _table(doc,
        headers=['Показатель', 'Значение'],
        rows=[
            ('Общая сумма поступлений (выручка)', _fmt(rev)),
            ('Количество оформленных заказов', str(orders_cnt)),
            ('Средняя стоимость заказа', _fmt(rev / orders_cnt if orders_cnt else 0)),
            ('Доходы наличными', _fmt(payments.get('cash', 0))),
            ('Доходы по банковской карте', _fmt(payments.get('card', 0))),
            ('Доходы банковским переводом', _fmt(payments.get('transfer', 0))),
        ],
        col_widths=[10, 7],
    )

    # ── Структура доходов по месяцам ──────────────────────────────────────────
    _h(doc, '3. СТРУКТУРА ДОХОДОВ ПО МЕСЯЦАМ', level=1)
    _p(doc,
       'В таблице ниже представлена разбивка совокупной выручки организации по '
       'календарным месяцам в пределах отчётного периода. Данная информация необходима '
       'для расчёта авансовых платежей по УСН нарастающим итогом.')

    if monthly:
        month_rows = [(m, _fmt(r), _fmt(r * Decimal('0.06'))) for m, r in monthly]
        _table(doc,
            headers=['Месяц', 'Выручка', 'Аванс УСН 6%'],
            rows=month_rows,
            col_widths=[5, 5.5, 5.5],
        )
    else:
        _p(doc, 'Данных за выбранный период не найдено.', align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # ── Доходы по видам услуг ─────────────────────────────────────────────────
    _h(doc, '4. ДОХОДЫ ПО ВИДАМ ОКАЗАННЫХ УСЛУГ', level=1)
    _p(doc,
       'Согласно данным информационной системы, за отчётный период были оказаны '
       'следующие виды услуг технического обслуживания и ремонта. Детализация по '
       'видам услуг применяется для подтверждения основного вида деятельности '
       'организации и расчёта статистической отчётности.')

    if services:
        svc_rows = [(n, c, _fmt(r)) for n, c, r in services]
        _table(doc,
            headers=['Наименование услуги', 'Кол-во', 'Выручка'],
            rows=svc_rows,
            col_widths=[8.5, 3, 5.5],
        )
    else:
        _p(doc, 'Данных об услугах за выбранный период не найдено.', align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # ── Расчёт налоговых обязательств ─────────────────────────────────────────
    _h(doc, '5. РАСЧЁТ НАЛОГОВЫХ ОБЯЗАТЕЛЬСТВ', level=1)
    _p(doc,
       'На основании результатов финансово-хозяйственной деятельности за отчётный '
       'период произведён расчёт налоговых обязательств организации по применяемой '
       'системе налогообложения.')

    _h(doc, '5.1. УСН «Доходы» (ставка 6%)', level=2)
    _p(doc, 'Организация применяет упрощённую систему налогообложения с объектом '
            'налогообложения «Доходы». Налоговая база равна сумме всех '
            'поступивших доходов без каких-либо вычетов.')
    _table(doc,
        headers=['Показатель', 'Сумма'],
        rows=[
            ('Доходы (налоговая база)', _fmt(rev)),
            ('Налоговая ставка УСН', '6%'),
            ('Исчисленный налог (6% × выручка)', _fmt(usn_tax)),
            ('Страховые взносы (не более 50% налога)', _fmt(usn_tax * Decimal('0.5'))),
            ('Налог к уплате (с учётом взносов)', _fmt(usn_tax * Decimal('0.5'))),
        ],
        col_widths=[10, 7],
    )

    _h(doc, '5.2. Справочно: УСН «Доходы минус расходы» (ставка 15%)', level=2)
    _p(doc, 'Для сравнительного анализа приведён расчёт налога по альтернативному '
            'объекту налогообложения. Расходами признана стоимость приобретённых '
            'запасных частей и расходных материалов, оприходованных на склад в '
            'течение отчётного периода.')
    _table(doc,
        headers=['Показатель', 'Сумма'],
        rows=[
            ('Доходы', _fmt(rev)),
            ('Расходы (себестоимость запчастей)', _fmt(parts_cost)),
            ('Налоговая база (Доходы − Расходы)', _fmt(rev - parts_cost)),
            ('Налоговая ставка УСН', '15%'),
            ('Исчисленный налог', _fmt(usn_tax_15)),
            ('Минимальный налог (1% от доходов)', _fmt(rev * Decimal('0.01'))),
        ],
        col_widths=[10, 7],
    )

    _h(doc, '5.3. Справочно: НДС (ставка 20%, если применимо)', level=2)
    _p(doc, 'Организация, применяющая УСН, не является плательщиком налога на '
            'добавленную стоимость (НДС) в соответствии со ст. 346.11 НК РФ. '
            'Ниже приведён расчёт для информационных целей на случай перехода '
            'на общую систему налогообложения.')
    _table(doc,
        headers=['Показатель', 'Сумма'],
        rows=[
            ('Доходы (с НДС)', _fmt(rev)),
            ('НДС (20/120 от суммы с НДС)', _fmt(vat_amount)),
            ('Доходы без НДС', _fmt(revenue_wo_vat)),
        ],
        col_widths=[10, 7],
    )

    # ── Заключение ────────────────────────────────────────────────────────────
    _h(doc, '6. ЗАКЛЮЧЕНИЕ', level=1)
    _p(doc,
       f'Настоящий отчёт составлен за период с {df} по {dt} на основании данных '
       f'информационной системы ООО «ВинАвто». Общая сумма задокументированных '
       f'доходов составила {_fmt(rev)}. '
       f'Налог к уплате по УСН «Доходы» (6%) составляет {_fmt(usn_tax * Decimal("0.5"))} '
       f'с учётом вычета уплаченных страховых взносов.')
    _p(doc,
       'Настоящий отчёт не заменяет официальную налоговую декларацию. Для подачи '
       'отчётности в налоговый орган необходимо воспользоваться формами, '
       'утверждёнными ФНС России.')
    _p(doc, ' ', size=8, space_after='0')

    _table(doc,
        headers=['Должность', 'ФИО', 'Подпись', 'Дата'],
        rows=[
            ('Руководитель', '___________________', '___________', gen_date),
            ('Главный бухгалтер', '___________________', '___________', gen_date),
        ],
        col_widths=[5, 5, 3.5, 3.5],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  WORD REPORT 2 — ОТЧЁТ О ПРОИЗВОДИТЕЛЬНОСТИ
# ══════════════════════════════════════════════════════════════════════════════

def build_word_productivity_report(data: dict) -> io.BytesIO:
    doc = Document()
    _setup_doc(doc)

    df    = _fmt_date(data['date_from'])
    dt    = _fmt_date(data['date_to'])
    total = data.get('orders_total', 0)
    done  = data.get('orders_completed', 0)
    rev   = data.get('total_revenue', Decimal('0'))
    avg   = data.get('avg_order_value', Decimal('0'))
    status_br  = data.get('orders_by_status', {})
    monthly    = data.get('monthly_orders', [])
    top_svc    = data.get('top_services', [])
    new_cl     = data.get('new_clients', 0)
    ret_cl     = data.get('returning_clients', 0)
    gen_date   = _fmt_date(date.today())

    pct_done = round(done / total * 100, 1) if total else 0

    # ── Заглавие ──────────────────────────────────────────────────────────────
    _p(doc, 'ООО «ВинАвто»', bold=True, center=True)
    _p(doc, 'ОТЧЁТ О ПРОИЗВОДИТЕЛЬНОСТИ И ЭФФЕКТИВНОСТИ ДЕЯТЕЛЬНОСТИ', bold=True, size=15, center=True)
    _p(doc, f'за период с {df} по {dt}', size=13, center=True)
    _p(doc, f'Дата составления: {gen_date}', size=12, center=True)
    _divider(doc)

    # ── 1. Общая сводка ───────────────────────────────────────────────────────
    _h(doc, '1. ОБЩАЯ СВОДКА ЗА ПЕРИОД', level=1)
    _p(doc,
       f'Настоящий отчёт отражает ключевые показатели эффективности работы '
       f'автосервиса ООО «ВинАвто» за период с {df} по {dt}. '
       f'В течение отчётного периода в производство было принято {total} заказов, '
       f'из которых закрыто (статус «Выполнен» или «Выдан») — {done} заказ(ов), '
       f'что составляет {pct_done}% от общего количества принятых заказов. '
       f'Совокупная выручка предприятия по закрытым и оплаченным заказам составила '
       f'{_fmt(rev)}, средняя стоимость одного заказа — {_fmt(avg)}.')

    _table(doc,
        headers=['Ключевой показатель', 'Значение'],
        rows=[
            ('Принято заказов всего', str(total)),
            ('Заказов завершено (Выполнен + Выдан)', str(done)),
            ('Процент закрытия заказов', f'{pct_done}%'),
            ('Новые клиенты за период', str(new_cl)),
            ('Постоянные клиенты (повторные обращения)', str(ret_cl)),
            ('Совокупная выручка', _fmt(rev)),
            ('Средняя стоимость заказа', _fmt(avg)),
        ],
        col_widths=[10, 7],
    )

    # ── 2. Динамика по месяцам ────────────────────────────────────────────────
    _h(doc, '2. ДИНАМИКА ЗАКАЗОВ И ВЫРУЧКИ ПО МЕСЯЦАМ', level=1)
    _p(doc,
       'Анализ динамики количества заказов и выручки в разрезе календарных месяцев '
       'позволяет выявить сезонные закономерности, периоды пиковой нагрузки и '
       'спада активности, что необходимо для планирования ресурсов предприятия, '
       'подбора и оптимизации численности персонала, а также формирования '
       'стратегии продаж и маркетинга на следующий период.')

    if monthly:
        m_rows = [(m, c, _fmt(r), _fmt(r / c if c else 0)) for m, c, r in monthly]
        _table(doc,
            headers=['Месяц', 'Заказов', 'Выручка', 'Средний чек'],
            rows=m_rows,
            col_widths=[4.5, 3, 5, 5],
        )
    else:
        _p(doc, 'Данных за выбранный период не найдено.')
        doc.add_paragraph()

    # ── 3. Анализ по статусам ─────────────────────────────────────────────────
    _h(doc, '3. АНАЛИЗ ЗАКАЗОВ ПО СТАТУСАМ', level=1)
    _p(doc,
       'Распределение заказов по статусам отражает текущее состояние производственного '
       'процесса. Высокая доля заказов в статусах «Ожидает» и «В ремонте» '
       'свидетельствует о загруженности мастерской. Значительное число заказов '
       'в статусе «Ожидание запчастей» указывает на необходимость оптимизации '
       'логистики поставок и расширения ассортимента складских запасов.')

    if status_br:
        total_st = sum(status_br.values()) or 1
        st_rows = [(lbl, cnt, f'{round(cnt / total_st * 100, 1)}%')
                   for lbl, cnt in status_br.items()]
        _table(doc,
            headers=['Статус заказа', 'Количество', 'Доля'],
            rows=st_rows,
            col_widths=[8, 3, 3],
        )
    else:
        _p(doc, 'Данных о статусах за выбранный период не найдено.')
        doc.add_paragraph()

    # ── 4. Топ услуг ──────────────────────────────────────────────────────────
    _h(doc, '4. РЕЙТИНГ ОКАЗЫВАЕМЫХ УСЛУГ', level=1)
    _p(doc,
       'Рейтинг оказываемых услуг формируется на основе статистики выполненных '
       'заказов за отчётный период. Наиболее востребованные услуги являются '
       'основой бизнеса предприятия и должны быть обеспечены достаточным '
       'количеством квалифицированных специалистов и расходных материалов. '
       'Услуги с низкой частотой заказов могут нуждаться в дополнительном '
       'продвижении или пересмотре ценовой политики.')

    if top_svc:
        svc_rows = [(n, c, _fmt(r), _fmt(r / c if c else 0)) for n, c, r in top_svc]
        _table(doc,
            headers=['Услуга', 'Заказов', 'Выручка', 'Средняя цена'],
            rows=svc_rows,
            col_widths=[7, 2.5, 4.5, 4],
        )
    else:
        _p(doc, 'Данных об услугах за выбранный период не найдено.')
        doc.add_paragraph()

    # ── 5. Клиентская база ────────────────────────────────────────────────────
    _h(doc, '5. АНАЛИЗ КЛИЕНТСКОЙ БАЗЫ', level=1)
    _p(doc,
       f'За отчётный период в организацию обратилось {new_cl} новых клиентов и '
       f'{ret_cl} клиентов с повторными обращениями. Доля постоянных клиентов '
       f'составляет {round(ret_cl / (new_cl + ret_cl) * 100) if (new_cl + ret_cl) else 0}% '
       f'от общего числа обращений, что является важным показателем лояльности '
       f'целевой аудитории. Повышение уровня удержания клиентов рекомендуется '
       f'достигать через программы лояльности, систему скидок для постоянных клиентов '
       f'и качественное постпродажное обслуживание.')

    # ── 6. Рекомендации ───────────────────────────────────────────────────────
    _h(doc, '6. РЕКОМЕНДАЦИИ', level=1)
    recommendations = [
        ('Оптимизация складских запасов', 'Сократить долю заказов в статусе «Ожидание запчастей» за счёт '
         'пересмотра минимальных остатков наиболее востребованных позиций на складе.'),
        ('Управление загрузкой персонала', 'Провести анализ распределения заказов между механиками; '
         'при необходимости ввести гибкий график и привлечь дополнительных специалистов.'),
        ('Развитие клиентской базы', 'Запустить реферальную программу для существующих клиентов '
         'с целью привлечения новых обращений.'),
        ('Ценовая политика', 'Пересмотреть прейскурант на услуги с учётом рыночных изменений '
         'и изменения себестоимости запасных частей.'),
    ]
    for title, text in recommendations:
        para = doc.add_paragraph(style='List Number')
        _spacing_xml(para)
        run = para.add_run(f'{title}. ')
        run.font.name = TNR
        run.font.size = Pt(14)
        run.font.bold = True
        _rfonts_xml(run)
        run2 = para.add_run(text)
        run2.font.name = TNR
        run2.font.size = Pt(14)
        _rfonts_xml(run2)

    doc.add_paragraph()
    _p(doc, ' ', size=4, space_after='0')
    _table(doc,
        headers=['Должность', 'ФИО', 'Подпись', 'Дата'],
        rows=[
            ('Руководитель', '___________________', '___________', gen_date),
        ],
        col_widths=[5, 5, 3.5, 3.5],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  WORD REPORT 3 — ОТЧЁТ ПО МЕХАНИКАМ
# ══════════════════════════════════════════════════════════════════════════════

def build_word_mechanics_report(data: dict) -> io.BytesIO:
    doc = Document()
    _setup_doc(doc)

    df       = _fmt_date(data['date_from'])
    dt       = _fmt_date(data['date_to'])
    mechanics = data.get('mechanics', [])
    gen_date  = _fmt_date(date.today())

    total_orders  = sum(m.get('orders_total', 0) for m in mechanics)
    total_revenue = sum(m.get('revenue', Decimal('0')) for m in mechanics)

    _p(doc, 'ООО «ВинАвто»', bold=True, center=True)
    _p(doc, 'ОТЧЁТ О ПРОИЗВОДИТЕЛЬНОСТИ МЕХАНИКОВ', bold=True, size=15, center=True)
    _p(doc, f'за период с {df} по {dt}', size=13, center=True)
    _p(doc, f'Дата составления: {gen_date}', size=12, center=True)
    _divider(doc)

    # ── 1. Введение ───────────────────────────────────────────────────────────
    _h(doc, '1. ВВЕДЕНИЕ', level=1)
    _p(doc,
       f'Настоящий отчёт содержит детализированные показатели эффективности '
       f'работы механиков ООО «ВинАвто» за период с {df} по {dt}. '
       f'В анализ включены: количество принятых и выполненных заказов, '
       f'суммарная выручка, сформированная каждым специалистом, средняя '
       f'стоимость одного заказа, а также рейтинг по ключевым показателям '
       f'эффективности (KPI). Данные используются руководством для принятия '
       f'управленческих решений в части оплаты труда, распределения нагрузки '
       f'и повышения квалификации персонала.')

    # ── 2. Сводка ─────────────────────────────────────────────────────────────
    _h(doc, '2. СВОДНЫЕ ПОКАЗАТЕЛИ КОМАНДЫ', level=1)
    _p(doc,
       f'В отчётном периоде в штате организации числится {len(mechanics)} '
       f'механик(ов). Суммарное количество заказов, выполненных командой, '
       f'составило {total_orders} единиц. Совокупная выручка, '
       f'сгенерированная механиками, равна {_fmt(total_revenue)}.')

    _table(doc,
        headers=['Показатель', 'Значение'],
        rows=[
            ('Количество механиков (активных)', str(len(mechanics))),
            ('Всего заказов за период', str(total_orders)),
            ('Суммарная выручка команды', _fmt(total_revenue)),
            ('Средняя выручка на 1 механика', _fmt(total_revenue / len(mechanics) if mechanics else 0)),
            ('Среднее число заказов на 1 механика', str(round(total_orders / len(mechanics), 1) if mechanics else 0)),
        ],
        col_widths=[10, 7],
    )

    # ── 3. Детальные показатели ───────────────────────────────────────────────
    _h(doc, '3. ДЕТАЛЬНЫЕ ПОКАЗАТЕЛИ ПО КАЖДОМУ МЕХАНИКУ', level=1)
    _p(doc,
       'Ниже представлены данные по каждому механику в разрезе ключевых '
       'показателей эффективности. Показатель «% закрытия» рассчитывается '
       'как отношение числа закрытых заказов к числу принятых в работу.')

    for m in mechanics:
        _h(doc, m['name'], level=2)
        mrev  = m.get('revenue', Decimal('0'))
        mtot  = m.get('orders_total', 0)
        mdone = m.get('orders_completed', 0)
        mavg  = mrev / mdone if mdone else Decimal('0')
        pct   = round(mdone / mtot * 100, 1) if mtot else 0
        _table(doc,
            headers=['Показатель', 'Значение'],
            rows=[
                ('Принято заказов', str(mtot)),
                ('Завершено заказов', str(mdone)),
                ('Процент закрытия', f'{pct}%'),
                ('Выручка по выполненным заказам', _fmt(mrev)),
                ('Средняя стоимость одного заказа', _fmt(mavg)),
            ],
            col_widths=[10, 7],
        )
        breakdown = m.get('status_breakdown', {})
        if breakdown:
            st_rows = [(lbl, cnt) for lbl, cnt in breakdown.items()]
            _table(doc,
                headers=['Статус заказов', 'Количество'],
                rows=st_rows,
                col_widths=[10, 7],
            )

    # ── 4. Рейтинг ────────────────────────────────────────────────────────────
    _h(doc, '4. РЕЙТИНГ МЕХАНИКОВ', level=1)
    _p(doc,
       'Рейтинг формируется по убыванию сгенерированной выручки. Данный показатель '
       'используется как один из ключевых критериев при расчёте премиальной части '
       'заработной платы согласно принятой в организации системе мотивации.')

    sorted_m = sorted(mechanics, key=lambda x: x.get('revenue', 0), reverse=True)
    rating_rows = []
    for i, m in enumerate(sorted_m, 1):
        mrev  = m.get('revenue', Decimal('0'))
        mtot  = m.get('orders_total', 0)
        mdone = m.get('orders_completed', 0)
        mavg  = mrev / mdone if mdone else Decimal('0')
        rating_rows.append((str(i), m['name'], str(mtot), str(mdone), _fmt(mrev), _fmt(mavg)))

    _table(doc,
        headers=['#', 'Механик', 'Принято', 'Закрыто', 'Выручка', 'Средний чек'],
        rows=rating_rows,
        col_widths=[1, 5.5, 2.5, 2.5, 4.5, 4],
    )

    # ── 5. Рекомендации ───────────────────────────────────────────────────────
    _h(doc, '5. РЕКОМЕНДАЦИИ ПО УПРАВЛЕНИЮ ПЕРСОНАЛОМ', level=1)
    _p(doc,
       'На основании анализа показателей за отчётный период рекомендуется '
       'провести следующие организационные мероприятия:')
    items = [
        'Для механиков с показателем закрытия заказов ниже 80% организовать '
        'дополнительное обучение или наставничество со стороны более опытных '
        'специалистов.',
        'Рассмотреть возможность перераспределения нагрузки в пользу '
        'механиков с максимальной производительностью для ускорения '
        'обработки заказов в периоды пиковой загруженности.',
        'Ввести систему ежемесячных KPI с привязкой премиальной части '
        'к показателям выручки и процента закрытия заказов.',
        'Провести квалификационную оценку специальностей исполнителей для '
        'формирования индивидуальных планов профессионального развития.',
    ]
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        _spacing_xml(para)
        run = para.add_run(item)
        run.font.name = TNR
        run.font.size = Pt(14)
        _rfonts_xml(run)

    doc.add_paragraph()
    _table(doc,
        headers=['Должность', 'ФИО', 'Подпись', 'Дата'],
        rows=[
            ('Руководитель', '___________________', '___________', gen_date),
            ('HR-менеджер / Старший механик', '___________________', '___________', gen_date),
        ],
        col_widths=[5.5, 4.5, 3.5, 3.5],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL REPORT 1 — ФИНАНСОВЫЙ ОТЧЁТ
# ══════════════════════════════════════════════════════════════════════════════

def build_excel_financial_report(data: dict) -> io.BytesIO:
    wb = openpyxl.Workbook()
    df = _fmt_date(data['date_from'])
    dt = _fmt_date(data['date_to'])

    # ── Sheet 1: Сводка ───────────────────────────────────────────────────────
    ws = wb.active
    ws.title = 'Сводка'
    ws.freeze_panes = 'A2'

    rev   = data.get('total_revenue', Decimal('0'))
    pcnt  = data.get('orders_count', 0)
    parts = data.get('parts_cost', Decimal('0'))
    pmeth = data.get('payment_by_method', {})

    title_font = XFont(name=TNR, size=14, bold=True, color='FFFFFF')
    ws['A1'] = f'Финансовый отчёт за период {df} — {dt}'
    ws['A1'].font   = title_font
    ws['A1'].fill   = _HEADER_FILL
    ws['A1'].alignment = XAlignment(horizontal='center')
    ws.merge_cells('A1:C1')

    _xh(ws, 2, ['Показатель', 'Значение', 'Примечание'])
    summary_rows = [
        ('Совокупная выручка', float(rev), ''),
        ('Себестоимость запчастей', float(parts), 'Из поставок'),
        ('Валовая прибыль', float(rev - parts), 'Выручка − Себестоимость'),
        ('Количество оплаченных заказов', pcnt, ''),
        ('Средний чек', float(rev / pcnt) if pcnt else 0, ''),
        ('Налог УСН 6%', float(rev * Decimal('0.06')), 'Расчётный'),
        ('Налог к уплате (после вычета взносов)', float(rev * Decimal('0.03')), 'Расчётный (−50%)'),
        ('Наличными', float(pmeth.get('cash', 0)), ''),
        ('Банковской картой', float(pmeth.get('card', 0)), ''),
        ('Банковским переводом', float(pmeth.get('transfer', 0)), ''),
    ]
    for r, row in enumerate(summary_rows, 3):
        _xr(ws, r, row, alt=(r % 2 == 0))
    _xcol_widths(ws, [40, 18, 30])
    ws.row_dimensions[1].height = 30

    # ── Sheet 2: По месяцам ───────────────────────────────────────────────────
    ws2 = wb.create_sheet('По месяцам')
    ws2.freeze_panes = 'A2'
    _xh(ws2, 1, ['Месяц', 'Выручка (₽)', 'Заказов', 'Средний чек (₽)', 'Налог УСН 6% (₽)'])
    monthly = data.get('monthly_data', [])
    for r, (month_name, revenue) in enumerate(monthly, 2):
        cnt  = data.get('monthly_orders_count', {}).get(month_name, 0)
        avg  = float(revenue / cnt) if cnt else 0.0
        tax  = float(revenue * Decimal('0.06'))
        _xr(ws2, r, [month_name, float(revenue), cnt, avg, tax], alt=(r % 2 == 0))
    _xcol_widths(ws2, [18, 20, 12, 20, 20])

    # ── Sheet 3: Методы оплаты ────────────────────────────────────────────────
    ws3 = wb.create_sheet('Методы оплаты')
    ws3.freeze_panes = 'A2'
    _xh(ws3, 1, ['Метод оплаты', 'Сумма (₽)', 'Доля (%)'])
    rev_f = float(rev) or 1.0
    methods = [
        ('Наличные',               float(pmeth.get('cash', 0))),
        ('Банковская карта',        float(pmeth.get('card', 0))),
        ('Банковский перевод',      float(pmeth.get('transfer', 0))),
    ]
    for r, (name, amount) in enumerate(methods, 2):
        pct = round(amount / rev_f * 100, 2)
        _xr(ws3, r, [name, amount, pct], alt=(r % 2 == 0))
    _xcol_widths(ws3, [30, 20, 15])

    # ── Sheet 4: По услугам ───────────────────────────────────────────────────
    ws4 = wb.create_sheet('По услугам')
    ws4.freeze_panes = 'A2'
    _xh(ws4, 1, ['Услуга', 'Кол-во заказов', 'Выручка (₽)', 'Средняя цена (₽)', 'Доля в выручке (%)'])
    services = data.get('services_revenue', [])
    for r, (name, cnt, service_rev) in enumerate(services, 2):
        avg_s = float(service_rev) / cnt if cnt else 0.0
        pct_s = round(float(service_rev) / rev_f * 100, 2)
        _xr(ws4, r, [name, cnt, float(service_rev), avg_s, pct_s], alt=(r % 2 == 0))
    _xcol_widths(ws4, [35, 16, 20, 20, 18])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL REPORT 2 — ОТЧЁТ ПО ЗАКАЗАМ
# ══════════════════════════════════════════════════════════════════════════════

def build_excel_orders_report(data: dict) -> io.BytesIO:
    wb  = openpyxl.Workbook()
    df  = _fmt_date(data['date_from'])
    dt  = _fmt_date(data['date_to'])

    ws = wb.active
    ws.title = 'Все заказы'
    ws.freeze_panes = 'A2'

    title_font = XFont(name=TNR, size=13, bold=True, color='FFFFFF')
    ws['A1'] = f'Реестр заказов: {df} — {dt}'
    ws['A1'].font = title_font
    ws['A1'].fill = _HEADER_FILL
    ws['A1'].alignment = XAlignment(horizontal='center')
    ws.merge_cells('A1:H1')

    _xh(ws, 2, ['ID', 'Дата', 'Клиент', 'Автомобиль', 'Механик', 'Статус', 'Сумма (₽)', 'Оплачено (₽)'])
    orders = data.get('orders_list', [])
    for r, o in enumerate(orders, 3):
        _xr(ws, r, [
            o.get('id', ''),
            o.get('date', ''),
            o.get('client', ''),
            o.get('car', ''),
            o.get('mechanic', '—'),
            o.get('status', ''),
            float(o.get('total_price', 0)),
            float(o.get('paid', 0)),
        ], alt=(r % 2 == 0))
    _xcol_widths(ws, [6, 14, 28, 22, 22, 18, 16, 16])
    ws.row_dimensions[1].height = 28

    # ── Sheet 2: По клиентам ──────────────────────────────────────────────────
    ws2 = wb.create_sheet('По клиентам')
    ws2.freeze_panes = 'A2'
    _xh(ws2, 1, ['Клиент', 'Телефон', 'Заказов', 'Сумма (₽)', 'Средний чек (₽)'])
    by_client = data.get('by_client', [])
    for r, c in enumerate(by_client, 2):
        _xr(ws2, r, [
            c.get('name', ''), c.get('phone', ''),
            c.get('count', 0), float(c.get('total', 0)),
            float(c.get('avg', 0)),
        ], alt=(r % 2 == 0))
    _xcol_widths(ws2, [30, 18, 12, 20, 20])

    # ── Sheet 3: По статусам ──────────────────────────────────────────────────
    ws3 = wb.create_sheet('По статусам')
    ws3.freeze_panes = 'A2'
    _xh(ws3, 1, ['Статус', 'Количество', 'Доля (%)'])
    status_br = data.get('orders_by_status', {})
    total_st  = sum(status_br.values()) or 1
    for r, (lbl, cnt) in enumerate(status_br.items(), 2):
        _xr(ws3, r, [lbl, cnt, round(cnt / total_st * 100, 2)], alt=(r % 2 == 0))
    _xcol_widths(ws3, [28, 14, 14])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL REPORT 3 — ОТЧЁТ ПО МЕХАНИКАМ
# ══════════════════════════════════════════════════════════════════════════════

def build_excel_mechanics_report(data: dict) -> io.BytesIO:
    wb = openpyxl.Workbook()
    df = _fmt_date(data['date_from'])
    dt = _fmt_date(data['date_to'])

    ws = wb.active
    ws.title = 'Производительность'
    ws.freeze_panes = 'A2'

    ws['A1'] = f'Производительность механиков: {df} — {dt}'
    ws['A1'].font  = XFont(name=TNR, size=13, bold=True, color='FFFFFF')
    ws['A1'].fill  = _HEADER_FILL
    ws['A1'].alignment = XAlignment(horizontal='center')
    ws.merge_cells('A1:G1')

    _xh(ws, 2, ['Механик', 'Принято заказов', 'Закрыто', '% закрытия',
                 'Выручка (₽)', 'Средний чек (₽)', 'Активных сейчас'])

    mechanics = data.get('mechanics', [])
    for r, m in enumerate(mechanics, 3):
        mtot  = m.get('orders_total', 0)
        mdone = m.get('orders_completed', 0)
        mrev  = float(m.get('revenue', 0))
        mavg  = mrev / mdone if mdone else 0.0
        pct   = round(mdone / mtot * 100, 1) if mtot else 0.0
        mact  = m.get('orders_active', 0)
        _xr(ws, r, [m['name'], mtot, mdone, pct, mrev, mavg, mact], alt=(r % 2 == 0))

    _xcol_widths(ws, [28, 16, 12, 12, 18, 18, 16])
    ws.row_dimensions[1].height = 28

    # ── Sheet 2: Детально по заказам ──────────────────────────────────────────
    ws2 = wb.create_sheet('Заказы по механикам')
    ws2.freeze_panes = 'A2'
    _xh(ws2, 1, ['Механик', 'ID заказа', 'Дата', 'Клиент', 'Статус', 'Сумма (₽)'])
    mechanic_orders = data.get('mechanic_orders', [])
    for r, o in enumerate(mechanic_orders, 2):
        _xr(ws2, r, [
            o.get('mechanic', ''),
            o.get('id', ''),
            o.get('date', ''),
            o.get('client', ''),
            o.get('status', ''),
            float(o.get('total_price', 0)),
        ], alt=(r % 2 == 0))
    _xcol_widths(ws2, [28, 10, 14, 28, 18, 16])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
